from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_times_new_roman(run) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def create_demo_beshtau_template(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.columns[0].width = Cm(8)
    header.columns[1].width = Cm(8)
    left = header.cell(0, 0)
    right = header.cell(0, 1)
    left.text = "ООО «Бештау Электроникс»\nТел.: +7 (000) 000-00-00\nwww.beshtau.ru\ninfo@beshtau.ru"
    right.text = "Исх. № {{outgoing_number}}\nот {{quote_date}}\n\n{{recipient_name}}\nИНН {{recipient_inn}}\n{{recipient_address}}"
    for p in right.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("{{intro_text}}")
    doc.add_paragraph("{{specification_text}}")

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.0), Cm(6.6), Cm(1.4), Cm(2.0), Cm(2.5), Cm(3.2)]
    headers = [
        "№ п/п",
        "Наименование товара",
        "Ед. изм.",
        "Общее количество",
        "Цена за единицу",
        "Общая сумма, руб., с учетом НДС",
    ]
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    markers = ["{{item_no}}", "{{item_name}}", "{{item_unit}}", "{{item_quantity}}", "{{item_unit_price}}", "{{item_line_total}}"]
    for index, text in enumerate(markers):
        table.cell(1, index).text = text
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    total = doc.add_table(rows=3, cols=2)
    total.style = "Table Grid"
    total.cell(0, 0).text = "ИТОГО, руб., с учетом НДС"
    total.cell(0, 1).text = "{{total_amount}}"
    total.cell(1, 0).text = "В том числе НДС {{vat_rate}}%"
    total.cell(1, 1).text = "{{vat_amount}}"
    total.cell(2, 0).text = "Сумма прописью"
    total.cell(2, 1).text = "{{total_amount_words}}"

    doc.add_paragraph("Срок поставки: {{delivery_term}}")
    doc.add_paragraph("Гарантия: {{warranty}}")
    doc.add_paragraph("Срок действия КП: {{valid_until}}")
    doc.add_paragraph("{{optional_conditions}}")

    sign = doc.add_paragraph()
    sign.add_run("{{signer_title}}").bold = True
    sign.add_run("\t\t\t")
    sign.add_run("{{signer_name}}").bold = True

    doc.save(str(path))


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        parts = text.split("\n")
        first.text = parts[0]
        for part in parts[1:]:
            first.add_break()
            first.add_text(part)
        set_times_new_roman(first)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        set_times_new_roman(run)


def set_cell_text(cell, text: str, alignment=None) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    set_paragraph_text(first, text)
    if alignment is not None:
        first.alignment = alignment
    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)


def normalize_paragraph_indent(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def copy_horizontal_indent(target, source) -> None:
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent


def prepare_beshtau_template_from_source(source_path: Path, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(source_path))

    if doc.tables:
        head = doc.tables[0]
        if len(head.rows) >= 1 and len(head.columns) >= 2:
            set_cell_text(head.cell(0, 0), "Исх.№ {{outgoing_number}} от {{quote_date}}")
            set_cell_text(
                head.cell(0, 1),
                "{{recipient_name}}\nИНН {{recipient_inn}}\n{{recipient_address}}",
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )

    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if non_empty:
        title = non_empty[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        normalize_paragraph_indent(title)
        for run in title.runs:
            set_times_new_roman(run)
            run.bold = True
            run.font.size = Pt(14)
    if len(non_empty) >= 2:
        set_paragraph_text(non_empty[1], "{{intro_text}}")
    if len(non_empty) >= 4:
        set_paragraph_text(
            non_empty[3],
            "Сумма: {{total_amount}} руб. ({{total_amount_words}}), в т.ч. НДС {{vat_rate}}% "
            "{{vat_amount}} руб. ({{vat_amount_words}})",
        )
    if len(non_empty) >= 5:
        set_paragraph_text(non_empty[4], "Срок поставки: {{delivery_term}}")
        copy_horizontal_indent(non_empty[3], non_empty[4])
        non_empty[3].paragraph_format.space_before = Pt(8)
        non_empty[3].paragraph_format.space_after = Pt(3)
    if len(non_empty) >= 6:
        set_paragraph_text(non_empty[5], "Срок гарантии: {{warranty}}")
        non_empty[4].paragraph_format.space_after = Pt(3)
    if len(non_empty) >= 7:
        set_paragraph_text(non_empty[6], "Коммерческое предложение действительно до {{valid_until}}")
        non_empty[5].paragraph_format.space_after = Pt(3)
        non_empty[6].paragraph_format.space_after = Pt(6)
    if len(non_empty) >= 8:
        set_paragraph_text(non_empty[7], "{{specification_text}}")
        optional = non_empty[7].insert_paragraph_before("{{optional_conditions}}")
        optional.style = non_empty[7].style

    if len(doc.tables) >= 2:
        goods = doc.tables[1]
        while len(goods.rows) > 3:
            goods._tbl.remove(goods.rows[2]._tr)
        marker_cells = goods.rows[1].cells
        markers = ["{{item_no}}", "{{item_name}}", "{{item_unit}}", "{{item_quantity}}", "{{item_unit_price}}", "{{item_line_total}}"]
        for idx, marker in enumerate(markers):
            set_cell_text(marker_cells[idx], marker)
        total_row = goods.rows[-1]
        for idx, cell in enumerate(total_row.cells):
            set_cell_text(cell, "ИТОГО" if idx < len(total_row.cells) - 1 else "{{total_amount}}")

    doc.save(str(output_path))

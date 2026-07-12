import shutil
import subprocess
import uuid
from copy import deepcopy
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell, _Row, Table
from docx.oxml.ns import qn
from lxml import etree

from app.config import get_settings
from app.models import Proposal
from app.services.money import format_money, format_numeric_date, format_ru_date, sanitize_filename_part


DEFAULT_SPECIFICATION_TEXT = (
    "\u0414\u0430\u043d\u043d\u043e\u0435 \u043a\u043e\u043c\u043c\u0435\u0440\u0447\u0435\u0441\u043a\u043e\u0435 \u043f\u0440\u0435\u0434\u043b\u043e\u0436\u0435\u043d\u0438\u0435 "
    "\u0434\u0435\u0439\u0441\u0442\u0432\u0438\u0442\u0435\u043b\u044c\u043d\u043e \u0434\u043b\u044f \u043a\u043e\u043d\u0444\u0438\u0433\u0443\u0440\u0430\u0446\u0438\u0438, "
    "\u043e\u043f\u0438\u0441\u0430\u043d\u043d\u043e\u0439 \u0432 \u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0438 "
    "\u201c\u0421\u043f\u0435\u0446\u0438\u0444\u0438\u043a\u0430\u0446\u0438\u044f \u21161\u201d, "
    "\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u043d\u043e\u043c \u043a \u043f\u0438\u0441\u044c\u043c\u0443."
)
OBJECT_DESCRIPTION_HEADING = "\u041e\u043f\u0438\u0441\u0430\u043d\u0438\u0435 \u043e\u0431\u044a\u0435\u043a\u0442\u0430 \u0437\u0430\u043a\u0443\u043f\u043a\u0438"


def _set_times_new_roman(run) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def default_intro_text(proposal: Proposal) -> str:
    if proposal.request_type.value == "with_request" and proposal.request_number and proposal.request_date:
        return (
            f"\u0418\u0437\u0443\u0447\u0438\u0432 \u043d\u0430\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u043d\u044b\u0439 \u0412\u0430\u043c\u0438 \u0437\u0430\u043f\u0440\u043e\u0441 \u2116{proposal.request_number} \u043e\u0442 "
            f"{format_numeric_date(proposal.request_date)} \u043e \u043f\u0440\u0435\u0434\u043e\u0441\u0442\u0430\u0432\u043b\u0435\u043d\u0438\u0438 \u0446\u0435\u043d\u043e\u0432\u043e\u0439 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u0438, "
            "\u043c\u044b, \u043d\u0438\u0436\u0435\u043f\u043e\u0434\u043f\u0438\u0441\u0430\u0432\u0448\u0438\u0435\u0441\u044f, \u043f\u0440\u0435\u0434\u043b\u0430\u0433\u0430\u0435\u043c \u043e\u0441\u0443\u0449\u0435\u0441\u0442\u0432\u0438\u0442\u044c "
            "\u043f\u043e\u0441\u0442\u0430\u0432\u043a\u0443 \u043e\u0431\u043e\u0440\u0443\u0434\u043e\u0432\u0430\u043d\u0438\u044f, \u0443\u043a\u0430\u0437\u0430\u043d\u043d\u043e\u0433\u043e \u0432 \u0437\u0430\u043f\u0440\u043e\u0441\u0435, "
            "\u043f\u043e\u0434\u0442\u0432\u0435\u0440\u0436\u0434\u0435\u043d\u043d\u0443\u044e \u043f\u0440\u0438\u043b\u0430\u0433\u0430\u0435\u043c\u043e\u0439 \u0442\u0430\u0431\u043b\u0438\u0446\u0435\u0439, \u0432 \u043a\u043e\u0442\u043e\u0440\u043e\u0439 "
            "\u0443\u043a\u0430\u0437\u0430\u043d\u0430 \u0446\u0435\u043d\u0430 \u0435\u0434\u0438\u043d\u0438\u0446\u044b \u0442\u043e\u0432\u0430\u0440\u0430 \u0438 \u043e\u0431\u0449\u0430\u044f \u0441\u0442\u043e\u0438\u043c\u043e\u0441\u0442\u044c:"
        )
    return "\u041f\u0440\u0435\u0434\u043b\u0430\u0433\u0430\u0435\u043c \u0440\u0430\u0441\u0441\u043c\u043e\u0442\u0440\u0435\u0442\u044c \u043a\u043e\u043c\u043c\u0435\u0440\u0447\u0435\u0441\u043a\u043e\u0435 \u043f\u0440\u0435\u0434\u043b\u043e\u0436\u0435\u043d\u0438\u0435 \u043d\u0430 \u043f\u043e\u0441\u0442\u0430\u0432\u043a\u0443 \u043e\u0431\u043e\u0440\u0443\u0434\u043e\u0432\u0430\u043d\u0438\u044f \u043d\u0430 \u0441\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0445 \u0443\u0441\u043b\u043e\u0432\u0438\u044f\u0445."


def legacy_intro_text(proposal: Proposal) -> str:
    if proposal.request_type.value == "with_request" and proposal.request_number and proposal.request_date:
        return (
            f"\u0418\u0437\u0443\u0447\u0438\u0432 \u043d\u0430\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u043d\u044b\u0439 \u0412\u0430\u043c\u0438 \u0437\u0430\u043f\u0440\u043e\u0441 \u2116{proposal.request_number} \u043e\u0442 "
            f"{format_ru_date(proposal.request_date)} \u043e \u043f\u0440\u0435\u0434\u043e\u0441\u0442\u0430\u0432\u043b\u0435\u043d\u0438\u0438 \u0446\u0435\u043d\u043e\u0432\u043e\u0439 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u0438, "
            "\u043c\u044b \u043f\u0440\u0435\u0434\u043b\u0430\u0433\u0430\u0435\u043c \u043f\u043e\u0441\u0442\u0430\u0432\u043a\u0443 \u043e\u0431\u043e\u0440\u0443\u0434\u043e\u0432\u0430\u043d\u0438\u044f \u043d\u0430 "
            "\u0441\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0445 \u0443\u0441\u043b\u043e\u0432\u0438\u044f\u0445."
        )
    return "\u041f\u0440\u0435\u0434\u043b\u0430\u0433\u0430\u0435\u043c \u0440\u0430\u0441\u0441\u043c\u043e\u0442\u0440\u0435\u0442\u044c \u043a\u043e\u043c\u043c\u0435\u0440\u0447\u0435\u0441\u043a\u043e\u0435 \u043f\u0440\u0435\u0434\u043b\u043e\u0436\u0435\u043d\u0438\u0435 \u043d\u0430 \u043f\u043e\u0441\u0442\u0430\u0432\u043a\u0443 \u043e\u0431\u043e\u0440\u0443\u0434\u043e\u0432\u0430\u043d\u0438\u044f \u043d\u0430 \u0441\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0445 \u0443\u0441\u043b\u043e\u0432\u0438\u044f\u0445."


def is_auto_intro_text(candidate: str | None, proposal: Proposal) -> bool:
    if candidate is None:
        return True
    normalized = " ".join(candidate.split())
    auto_variants = {
        " ".join(default_intro_text(proposal).split()),
        " ".join(legacy_intro_text(proposal).split()),
    }
    return normalized == "" or normalized in auto_variants


def proposal_context(proposal: Proposal) -> dict[str, str]:
    recipient = proposal.recipient_name.upper() if proposal.recipient_uppercase else proposal.recipient_name
    signer_title = proposal.signer.title if proposal.signer else "\u0413\u0435\u043d\u0435\u0440\u0430\u043b\u044c\u043d\u044b\u0439 \u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440"
    signer_name = proposal.signer.name if proposal.signer else "\u0412.\u041e. \u0413\u0430\u043b\u0443\u0441\u0442\u044f\u043d"
    delivery_unit = "\u0440\u0430\u0431\u043e\u0447\u0438\u0445 \u0434\u043d\u0435\u0439" if proposal.delivery_term_unit == "working_days" else "\u043a\u0430\u043b\u0435\u043d\u0434\u0430\u0440\u043d\u044b\u0445 \u0434\u043d\u0435\u0439"
    delivery_term = f"{proposal.delivery_term_value} {delivery_unit}" if proposal.delivery_term_value else ""
    optional_conditions = []
    if proposal.payment_terms:
        optional_conditions.append(f"\u0423\u0441\u043b\u043e\u0432\u0438\u044f \u043e\u043f\u043b\u0430\u0442\u044b: {proposal.payment_terms}")
    if proposal.delivery_terms:
        optional_conditions.append(f"\u0423\u0441\u043b\u043e\u0432\u0438\u044f \u0434\u043e\u0441\u0442\u0430\u0432\u043a\u0438: {proposal.delivery_terms}")
    if proposal.delivery_place:
        optional_conditions.append(f"\u041c\u0435\u0441\u0442\u043e \u043f\u043e\u0441\u0442\u0430\u0432\u043a\u0438: {proposal.delivery_place}")
    if proposal.request_type.value == "with_request" and proposal.request_number and proposal.request_date:
        nitrino_intro = (
            f"В ответ на Ваш запрос №{proposal.request_number} от {format_numeric_date(proposal.request_date)} г. "
            "предлагаем Вам рассмотреть к поставке товар, соответствующий Вашим техническим требованиям:"
        )
    else:
        nitrino_intro = "Предлагаем Вам рассмотреть к поставке товар, соответствующий Вашим техническим требованиям:"

    return {
        "recipient_name": recipient,
        "recipient_inn": proposal.recipient_inn or "",
        "recipient_email": proposal.recipient_email or "",
        "recipient_address": proposal.recipient_address or "",
        "quote_date": format_ru_date(proposal.quote_date),
        "quote_date_short": proposal.quote_date.strftime("%d.%m.%y"),
        "quote_date_numeric": format_numeric_date(proposal.quote_date),
        "outgoing_number": proposal.outgoing_number,
        "intro_text": proposal.intro_text or default_intro_text(proposal),
        "specification_text": proposal.specification_text,
        "delivery_term": delivery_term,
        "warranty": f"{proposal.warranty_months} \u043c\u0435\u0441.",
        "valid_until": format_ru_date(proposal.valid_until),
        "payment_terms": proposal.payment_terms or "",
        "delivery_terms": proposal.delivery_terms or "",
        "delivery_place": proposal.delivery_place or "",
        "optional_conditions": "\n".join(optional_conditions),
        "total_amount": format_money(proposal.total_amount),
        "vat_rate": format_money(proposal.vat_rate).rstrip("0").rstrip(","),
        "vat_amount": format_money(proposal.vat_amount),
        "total_amount_words": proposal.total_amount_words,
        "vat_amount_words": proposal.vat_amount_words,
        "signer_title": signer_title,
        "signer_name": signer_name,
        "kartas_conditions": " ".join(
            part
            for part in [
                f"Ориентировочный срок производства {delivery_term}." if delivery_term else "",
                f"Условия оплаты: {proposal.payment_terms}." if proposal.payment_terms else "",
                f"Условия доставки: {proposal.delivery_terms}." if proposal.delivery_terms else "",
                f"Место поставки: {proposal.delivery_place}." if proposal.delivery_place else "",
                f"Срок гарантии {proposal.warranty_months} мес.",
            ]
            if part
        ),
        "nitrino_intro": nitrino_intro,
    }


def _replace_in_paragraph(paragraph, context: dict[str, str]) -> None:
    for run in paragraph.runs:
        text = run.text
        changed = False
        for key, value in context.items():
            placeholder = "{{" + key + "}}"
            if placeholder in text:
                text = text.replace(placeholder, value)
                changed = True
        if changed:
            if "\n" not in text:
                run.text = text
            else:
                parts = text.split("\n")
                run.text = parts[0]
                for part in parts[1:]:
                    run.add_break()
                    run.add_text(part)
            _set_times_new_roman(run)


def _replace_in_cell(cell: _Cell, context: dict[str, str]) -> None:
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, context)
    for table in cell.tables:
        _replace_in_table(table, context)


def _replace_in_table(table: Table, context: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, context)


def _replace_everywhere(doc: Document, context: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)
    for table in doc.tables:
        _replace_in_table(table, context)
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            _replace_in_paragraph(paragraph, context)
        for table in section.header.tables + section.footer.tables:
            _replace_in_table(table, context)


def _paragraph_has_only_placeholder(paragraph, key: str) -> bool:
    text = paragraph.text.strip()
    return text == "{{" + key + "}}"


def _paragraph_contains_placeholder(paragraph, key: str) -> bool:
    return "{{" + key + "}}" in paragraph.text


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _remove_table_row(row: _Row) -> None:
    row._tr.getparent().remove(row._tr)


def _remove_object_description_heading(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if OBJECT_DESCRIPTION_HEADING in paragraph.text:
            _remove_paragraph(paragraph)
    for table in doc.tables:
        for row in list(table.rows):
            if any(OBJECT_DESCRIPTION_HEADING in cell.text for cell in row.cells):
                _remove_table_row(row)


def _remove_empty_block_placeholders(doc: Document, context: dict[str, str]) -> None:
    removable_keys = ["specification_text", "optional_conditions"]
    for key in removable_keys:
        if context.get(key, "") != "":
            continue
        for paragraph in list(doc.paragraphs):
            if _paragraph_has_only_placeholder(paragraph, key):
                _remove_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        if _paragraph_has_only_placeholder(paragraph, key):
                            _remove_paragraph(paragraph)
    if context.get("delivery_term", "") == "":
        for paragraph in list(doc.paragraphs):
            if _paragraph_contains_placeholder(paragraph, "delivery_term"):
                _remove_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        if _paragraph_contains_placeholder(paragraph, "delivery_term"):
                            _remove_paragraph(paragraph)


def _force_body_times_new_roman(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_times_new_roman(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_times_new_roman(run)


def _has_specification_page(proposal: Proposal) -> bool:
    return len((proposal.specification_text or "").strip()) > 10


def _specification_page_body(proposal: Proposal) -> str:
    text = (proposal.specification_text or "").strip()
    if " ".join(text.split()) == " ".join(DEFAULT_SPECIFICATION_TEXT.split()):
        return ""
    return text


def _add_text_with_breaks(paragraph, text: str) -> None:
    lines = text.splitlines() or [""]
    run = paragraph.add_run(lines[0])
    _set_times_new_roman(run)
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _append_specification_page(doc: Document, proposal: Proposal) -> None:
    template_version = getattr(proposal, "template_version", None)
    if (
        getattr(template_version, "placeholder_schema", None) in {"builtin-kartas-v1", "builtin-nitrino-v1"}
        and " ".join((proposal.specification_text or "").split()) == " ".join(DEFAULT_SPECIFICATION_TEXT.split())
    ):
        return
    if not _has_specification_page(proposal):
        return
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    title = f"Спецификация №1 к коммерческому предложению Исх.№ {proposal.outgoing_number} от {format_ru_date(proposal.quote_date)}"
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_text_with_breaks(title_paragraph, title)

    specification_text = _specification_page_body(proposal)
    if specification_text:
        doc.add_paragraph()
        specification_paragraph = doc.add_paragraph()
        _add_text_with_breaks(specification_paragraph, specification_text)


def _row_has_item_placeholders(row: _Row) -> bool:
    return any("{{item_" in cell.text for cell in row.cells)


def _row_context(item, index: int) -> dict[str, str]:
    return {
        "item_no": str(index),
        "item_name": item.display_name or item.name,
        "item_unit": item.unit,
        "item_quantity": str(item.quantity),
        "item_unit_price": format_money(item.unit_price_vat),
        "item_line_total": format_money(item.line_total),
    }


def _replace_items_table(doc: Document, proposal: Proposal) -> None:
    for table in doc.tables:
        marker_index = next((idx for idx, row in enumerate(table.rows) if _row_has_item_placeholders(row)), None)
        if marker_index is None:
            continue
        marker_row = table.rows[marker_index]
        template_tr = deepcopy(marker_row._tr)

        for index, item in enumerate(sorted(proposal.items, key=lambda x: x.sort_order), start=1):
            new_tr = deepcopy(template_tr)
            marker_row._tr.addprevious(new_tr)
            row = _Row(new_tr, table)
            for cell in row.cells:
                _replace_in_cell(cell, _row_context(item, index))
        marker_row._tr.getparent().remove(marker_row._tr)
        return


def _replace_xml_placeholders(element, context: dict[str, str]) -> None:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    word_namespace = namespace["w"]
    for text_node in list(element.xpath(".//w:t", namespaces=namespace)):
        text = text_node.text or ""
        for key, value in context.items():
            text = text.replace("{{" + key + "}}", value)
        lines = text.replace("\r\n", "\n").split("\n")
        text_node.text = lines[0]
        if len(lines) > 1:
            run = text_node.getparent()
            insert_at = run.index(text_node) + 1
            for line in lines[1:]:
                run.insert(insert_at, etree.Element(f"{{{word_namespace}}}br"))
                insert_at += 1
                continuation = etree.Element(f"{{{word_namespace}}}t")
                continuation.text = line
                run.insert(insert_at, continuation)
                insert_at += 1


def _set_xml_text(element, value: str, namespace: dict[str, str], collapse_paragraphs: bool = False) -> None:
    text_nodes = element.xpath(".//w:t", namespaces=namespace)
    if not text_nodes:
        raise ValueError("В фирменном шаблоне не найден ожидаемый текстовый узел")
    text_nodes[0].text = value.splitlines()[0] if value.splitlines() else ""
    for node in text_nodes[1:]:
        node.text = ""
    for break_or_tab in element.xpath(".//w:br | .//w:tab", namespaces=namespace):
        break_or_tab.getparent().remove(break_or_tab)
    lines = value.replace("\r\n", "\n").split("\n")
    if len(lines) > 1:
        run = text_nodes[0].getparent()
        insert_at = run.index(text_nodes[0]) + 1
        word_namespace = namespace["w"]
        for line in lines[1:]:
            run.insert(insert_at, etree.Element(f"{{{word_namespace}}}br"))
            insert_at += 1
            continuation = etree.Element(f"{{{word_namespace}}}t")
            continuation.text = line
            run.insert(insert_at, continuation)
            insert_at += 1
    if collapse_paragraphs:
        for extra_paragraph in element.findall("w:p", namespace)[1:]:
            element.remove(extra_paragraph)


def _render_branded_ooxml(proposal: Proposal, template_path: Path, output_path: Path) -> None:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(template_path, "r") as source_package:
        root = etree.fromstring(source_package.read("word/document.xml"))
        marker_row = next(
            (
                row
                for row in root.xpath(".//w:tr", namespaces=namespace)
                if "{{item_" in "".join(row.xpath(".//w:t/text()", namespaces=namespace))
            ),
            None,
        )
        if marker_row is None:
            raise ValueError("В фирменном шаблоне не найдена строка товаров")

        parent = marker_row.getparent()
        insert_at = parent.index(marker_row)
        for index, item in enumerate(sorted(proposal.items, key=lambda value: value.sort_order), start=1):
            row = deepcopy(marker_row)
            _replace_xml_placeholders(row, _row_context(item, index))
            parent.insert(insert_at, row)
            insert_at += 1
        parent.remove(marker_row)
        _replace_xml_placeholders(root, proposal_context(proposal))

        patched_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)
        temporary_path = output_path.with_suffix(output_path.suffix + ".tmp")
        with ZipFile(temporary_path, "w") as output_package:
            for info in source_package.infolist():
                content = patched_xml if info.filename == "word/document.xml" else source_package.read(info.filename)
                output_package.writestr(info, content)
    temporary_path.replace(output_path)


def _render_nitrino_ooxml(proposal: Proposal, template_path: Path, output_path: Path) -> None:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    context = proposal_context(proposal)
    with ZipFile(template_path, "r") as source_package:
        root = etree.fromstring(source_package.read("word/document.xml"))
        body = root.find("w:body", namespace)
        if body is None:
            raise ValueError("В шаблоне НИТРИНО отсутствует тело документа")
        paragraphs = body.findall("w:p", namespace)
        tables = body.findall("w:tbl", namespace)
        if len(paragraphs) < 4 or len(tables) < 3:
            raise ValueError("Структура шаблона НИТРИНО не соответствует ожидаемой")

        metadata_cells = tables[0].findall("w:tr", namespace)[0].findall("w:tc", namespace)
        _set_xml_text(
            metadata_cells[0],
            f"Исх: № {proposal.outgoing_number} от {context['quote_date_numeric']} г.",
            namespace,
        )
        _set_xml_text(metadata_cells[1], context["recipient_name"], namespace)
        _set_xml_text(paragraphs[3], context["nitrino_intro"], namespace)

        goods_rows = tables[1].findall("w:tr", namespace)
        if len(goods_rows) < 5 or len(goods_rows[1].findall("w:tc", namespace)) != 6:
            raise ValueError("Таблица товаров НИТРИНО не соответствует ожидаемой геометрии")
        source_item_rows = goods_rows[1:4]
        total_row = goods_rows[4]
        insert_at = tables[1].index(total_row)
        for index, item in enumerate(sorted(proposal.items, key=lambda value: value.sort_order), start=1):
            if index <= len(source_item_rows):
                row = source_item_rows[index - 1]
            else:
                row = deepcopy(source_item_rows[-1])
                for element in row.iter():
                    for attribute in list(element.attrib):
                        if attribute.endswith("}paraId") or attribute.endswith("}textId"):
                            del element.attrib[attribute]
                tables[1].insert(insert_at, row)
                insert_at += 1
            row_context = _row_context(item, index)
            values = [
                row_context["item_no"],
                row_context["item_name"],
                row_context["item_unit"],
                row_context["item_quantity"],
                row_context["item_unit_price"],
                row_context["item_line_total"],
            ]
            for cell, value in zip(row.findall("w:tc", namespace), values):
                _set_xml_text(cell, value, namespace, collapse_paragraphs=True)
        for unused_row in source_item_rows[len(proposal.items) :]:
            tables[1].remove(unused_row)
        _set_xml_text(total_row.findall("w:tc", namespace)[1], context["total_amount"], namespace)

        summary_rows = tables[2].findall("w:tr", namespace)
        total_summary = (
            f"{context['total_amount']} руб., {context['total_amount_words']}, включая НДС "
            f"({context['vat_rate']}%) в сумме {context['vat_amount']} руб., {context['vat_amount_words']}"
        )
        _set_xml_text(summary_rows[0].findall("w:tc", namespace)[1], total_summary, namespace)
        _set_xml_text(
            summary_rows[2].findall("w:tc", namespace)[1],
            f"Коммерческое предложение действительно до {context['valid_until']}",
            namespace,
        )
        _set_xml_text(
            summary_rows[4].findall("w:tc", namespace)[1],
            f"{context['signer_title']} {context['signer_name']}___________ м.п.",
            namespace,
        )

        patched_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)
        temporary_path = output_path.with_suffix(output_path.suffix + ".tmp")
        with ZipFile(temporary_path, "w") as output_package:
            for info in source_package.infolist():
                content = patched_xml if info.filename == "word/document.xml" else source_package.read(info.filename)
                output_package.writestr(info, content)
    temporary_path.replace(output_path)


def _output_stem(proposal: Proposal, suffix: str = "") -> str:
    proposal_date = proposal.quote_date.strftime("%d.%m.%Y")
    template = getattr(getattr(proposal, "template_version", None), "template", None)
    organization = getattr(template, "organization", None) or "Бештау"
    parts = [
        f"КП {organization}",
        sanitize_filename_part(proposal.recipient_inn or "\u0431\u0435\u0437 \u0418\u041d\u041d", max_len=20),
        sanitize_filename_part(proposal.recipient_name, max_len=80),
        proposal_date,
    ]
    return f"{' - '.join(parts)}{suffix}"


def _make_output_paths(proposal: Proposal, suffix: str = "") -> tuple[Path, Path]:
    settings = get_settings()
    stem = _output_stem(proposal, suffix=suffix)
    proposal_dir = settings.generated_dir / str(proposal.id)
    proposal_dir.mkdir(parents=True, exist_ok=True)
    return proposal_dir / f"{stem}.docx", proposal_dir / f"{stem}.pdf"


def render_docx(proposal: Proposal, template_path: Path, preview: bool = False) -> Path:
    template_schema = getattr(getattr(proposal, "template_version", None), "placeholder_schema", None)
    if template_schema in {"builtin-kartas-v1", "builtin-nitrino-v1"}:
        suffix = f"_preview_{uuid.uuid4().hex[:8]}" if preview else ""
        docx_path, _ = _make_output_paths(proposal, suffix=suffix)
        if preview:
            docx_path = get_settings().previews_dir / docx_path.name
        if template_schema == "builtin-nitrino-v1":
            _render_nitrino_ooxml(proposal, template_path, docx_path)
        else:
            _render_branded_ooxml(proposal, template_path, docx_path)
        if _has_specification_page(proposal) and _specification_page_body(proposal):
            doc = Document(str(docx_path))
            _append_specification_page(doc, proposal)
            doc.save(str(docx_path))
        return docx_path

    doc = Document(str(template_path))
    context = proposal_context(proposal)
    _replace_items_table(doc, proposal)
    _remove_object_description_heading(doc)
    _remove_empty_block_placeholders(doc, context)
    _replace_everywhere(doc, context)
    _append_specification_page(doc, proposal)
    _force_body_times_new_roman(doc)
    suffix = f"_preview_{uuid.uuid4().hex[:8]}" if preview else ""
    docx_path, _ = _make_output_paths(proposal, suffix=suffix)
    if preview:
        docx_path = get_settings().previews_dir / docx_path.name
    doc.save(str(docx_path))
    return docx_path


def convert_docx_to_pdf(docx_path: Path, output_dir: Path | None = None) -> Path:
    settings = get_settings()
    soffice = settings.libreoffice_bin
    if shutil.which(soffice) is None and not Path(soffice).exists():
        raise RuntimeError("LibreOffice не найден. Укажите LIBREOFFICE_BIN или установите libreoffice в контейнер.")
    out_dir = output_dir or docx_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=90,
    )
    pdf_path = out_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf_path.exists():
        raise RuntimeError(f"Не удалось сконвертировать DOCX в PDF: {result.stderr or result.stdout}")
    return pdf_path


def generate_files(proposal: Proposal, template_path: Path) -> tuple[Path, Path]:
    docx_path, pdf_path = _make_output_paths(proposal)
    rendered_docx = render_docx(proposal, template_path, preview=False)
    if rendered_docx != docx_path:
        shutil.move(str(rendered_docx), str(docx_path))
    converted_pdf = convert_docx_to_pdf(docx_path, output_dir=docx_path.parent)
    if converted_pdf != pdf_path:
        shutil.move(str(converted_pdf), str(pdf_path))
    return docx_path, pdf_path


def generate_preview(proposal: Proposal, template_path: Path) -> Path:
    docx_path = render_docx(proposal, template_path, preview=True)
    return convert_docx_to_pdf(docx_path, output_dir=get_settings().previews_dir)




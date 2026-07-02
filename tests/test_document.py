import unittest
from datetime import date
from types import SimpleNamespace

from docx import Document

from app.services.document import _output_stem, _remove_object_description_heading, proposal_context


class DocumentNamingTest(unittest.TestCase):
    def test_output_stem_uses_brand_inn_recipient_and_date(self):
        proposal = SimpleNamespace(
            quote_date=date(2026, 6, 23),
            recipient_inn="7727482783",
            recipient_name="\u041e\u041e\u041e \u0420\u043e\u043c\u0430\u0448\u043a\u0430",
        )

        self.assertEqual(_output_stem(proposal), "\u041a\u041f \u0411\u0435\u0448\u0442\u0430\u0443 - 7727482783 - \u041e\u041e\u041e_\u0420\u043e\u043c\u0430\u0448\u043a\u0430 - 23.06.2026")
        self.assertEqual(_output_stem(proposal, suffix="_preview_1234"), "\u041a\u041f \u0411\u0435\u0448\u0442\u0430\u0443 - 7727482783 - \u041e\u041e\u041e_\u0420\u043e\u043c\u0430\u0448\u043a\u0430 - 23.06.2026_preview_1234")


class DocumentContentTest(unittest.TestCase):
    def test_context_uses_selected_signer(self):
        proposal = SimpleNamespace(
            signer=SimpleNamespace(title="Исполнительный директор", name="И.И. Иванов"),
            recipient_name="ООО Ромашка",
            recipient_uppercase=False,
            recipient_inn=None,
            recipient_email=None,
            recipient_address=None,
            quote_date=date(2026, 6, 23),
            outgoing_number="2306/1/М",
            intro_text="",
            specification_text="",
            delivery_term_unit="working_days",
            delivery_term_value=None,
            warranty_months=12,
            valid_until=date(2026, 7, 23),
            payment_terms=None,
            delivery_terms=None,
            delivery_place=None,
            total_amount=0,
            vat_rate=22,
            vat_amount=0,
            total_amount_words="",
            vat_amount_words="",
            request_type=SimpleNamespace(value="without_request"),
            request_number=None,
            request_date=None,
        )

        context = proposal_context(proposal)

        self.assertEqual(context["signer_title"], "Исполнительный директор")
        self.assertEqual(context["signer_name"], "И.И. Иванов")

    def test_removes_object_description_heading(self):
        doc = Document()
        doc.add_paragraph("Вступление")
        doc.add_paragraph("Описание объекта закупки")
        doc.add_paragraph("Финальный текст")

        _remove_object_description_heading(doc)

        self.assertEqual([paragraph.text for paragraph in doc.paragraphs], ["Вступление", "Финальный текст"])


if __name__ == "__main__":
    unittest.main()

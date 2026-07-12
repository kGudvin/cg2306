import hashlib
import tempfile
import unittest
from datetime import date
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT

from app.services.document import DEFAULT_SPECIFICATION_TEXT, _render_nitrino_ooxml
from app.template_seed import prepare_nitrino_template_from_source


SOURCE = Path(__file__).resolve().parents[1] / "НИТРИНО ШАБЛОН.docx"


def package_hashes(path: Path) -> dict[str, str]:
    with ZipFile(path) as package:
        return {info.filename: hashlib.sha256(package.read(info.filename)).hexdigest() for info in package.infolist()}


class NitrinoTemplateTest(unittest.TestCase):
    @staticmethod
    def sample_proposal():
        item = SimpleNamespace(
            sort_order=1,
            display_name="Компьютер NITRINOnet S600\nРеестровый номер 10577232\nСтрана происхождения — Российская Федерация",
            name="Компьютер NITRINOnet S600",
            unit="Шт.",
            quantity=3,
            unit_price_vat=Decimal("110400"),
            line_total=Decimal("331200"),
        )
        return SimpleNamespace(
            items=[item],
            signer=SimpleNamespace(title="Директор НЬЮ АЙ ТИ", name="Тюрин Д. А."),
            recipient_name="ГБУ «Заказчик»",
            recipient_uppercase=False,
            recipient_inn="5610127776",
            recipient_email="ignored@example.com",
            recipient_address="Этот адрес не должен выводиться",
            quote_date=date(2026, 7, 12),
            outgoing_number="1207/26/М",
            intro_text="",
            specification_text=DEFAULT_SPECIFICATION_TEXT,
            delivery_term_unit="working_days",
            delivery_term_value=30,
            warranty_months=12,
            valid_until=date(2026, 8, 12),
            payment_terms=None,
            delivery_terms=None,
            delivery_place=None,
            total_amount=Decimal("331200"),
            vat_rate=Decimal("22"),
            vat_amount=Decimal("59724.59"),
            total_amount_words="триста тридцать одна тысяча двести рублей 00 копеек",
            vat_amount_words="пятьдесят девять тысяч семьсот двадцать четыре рубля 59 копеек",
            request_type=SimpleNamespace(value="with_request"),
            request_number="42",
            request_date=date(2026, 7, 10),
            template_version=SimpleNamespace(placeholder_schema="builtin-nitrino-v1"),
        )

    def test_prepares_nitrino_as_byte_exact_source_copy(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "prepared.docx"
            prepare_nitrino_template_from_source(SOURCE, output)

            doc = Document(output)
            self.assertEqual(doc.sections[0].orientation, WD_ORIENT.PORTRAIT)
            self.assertEqual(hashlib.sha256(SOURCE.read_bytes()).hexdigest(), hashlib.sha256(output.read_bytes()).hexdigest())

    def test_final_render_repeats_items_and_omits_recipient_email_address(self):
        with tempfile.TemporaryDirectory() as tmp:
            prepared = Path(tmp) / "prepared.docx"
            output = Path(tmp) / "output.docx"
            prepare_nitrino_template_from_source(SOURCE, prepared)
            _render_nitrino_ooxml(self.sample_proposal(), prepared, output)

            doc = Document(output)
            all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            all_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
            self.assertIn("ГБУ «Заказчик»", all_text)
            self.assertIn("Компьютер NITRINOnet S600", all_text)
            self.assertIn("Реестровый номер 10577232", all_text)
            self.assertNotIn("ignored@example.com", all_text)
            self.assertNotIn("Этот адрес не должен выводиться", all_text)
            self.assertEqual({name for name in package_hashes(prepared) if package_hashes(prepared)[name] != package_hashes(output)[name]}, {"word/document.xml"})


if __name__ == "__main__":
    unittest.main()

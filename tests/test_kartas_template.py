import hashlib
import tempfile
import unittest
from datetime import date
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT

from app.template_seed import prepare_kartas_template_from_source
from app.services.document import DEFAULT_SPECIFICATION_TEXT, _append_specification_page, _render_kartas_ooxml


SOURCE = Path(__file__).resolve().parents[1] / "КП КАРТАС ШАБЛОН.docx"


def part_hash(path: Path, part_name: str) -> str:
    with ZipFile(path) as package:
        return hashlib.sha256(package.read(part_name)).hexdigest()


class KartasTemplateTest(unittest.TestCase):
    @staticmethod
    def sample_proposal():
        item = SimpleNamespace(
            sort_order=1,
            display_name="Сервер KCG Rack 2U",
            name="Сервер KCG Rack 2U",
            unit="шт.",
            quantity=2,
            unit_price_vat=Decimal("245000"),
            line_total=Decimal("490000"),
        )
        return SimpleNamespace(
            items=[item],
            signer=None,
            recipient_name="ООО «ПС Парт»",
            recipient_uppercase=False,
            recipient_inn="7707083893",
            recipient_email=None,
            recipient_address="г. Москва",
            quote_date=date(2026, 7, 12),
            outgoing_number="1207/39/М",
            intro_text="",
            specification_text=DEFAULT_SPECIFICATION_TEXT,
            delivery_term_unit="working_days",
            delivery_term_value=30,
            warranty_months=12,
            valid_until=date(2026, 8, 12),
            payment_terms="100% предоплата",
            delivery_terms=None,
            delivery_place=None,
            total_amount=Decimal("490000"),
            vat_rate=Decimal("22"),
            vat_amount=Decimal("88360.66"),
            total_amount_words="",
            vat_amount_words="",
            request_type=SimpleNamespace(value="without_request"),
            request_number=None,
            request_date=None,
            template_version=SimpleNamespace(placeholder_schema="builtin-kartas-v1"),
        )

    def test_default_specification_does_not_add_an_empty_page(self):
        doc = Document()
        proposal = SimpleNamespace(
            specification_text=DEFAULT_SPECIFICATION_TEXT,
            template_version=SimpleNamespace(placeholder_schema="builtin-kartas-v1"),
        )

        _append_specification_page(doc, proposal)

        self.assertEqual(len(doc.sections), 1)

    def test_prepares_placeholders_and_preserves_branded_assets(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "kartas_prepared.docx"
            prepare_kartas_template_from_source(SOURCE, output)

            doc = Document(output)
            self.assertEqual(doc.sections[0].orientation, WD_ORIENT.LANDSCAPE)
            self.assertEqual(doc.paragraphs[1].text, "Исх. {{outgoing_number}}")
            self.assertEqual(doc.paragraphs[2].text, "От {{quote_date_short}}")
            self.assertEqual(doc.paragraphs[5].text, "{{recipient_name}}")
            self.assertIn("{{kartas_conditions}}", doc.paragraphs[6].text)

            marker_row = doc.tables[1].rows[1]
            self.assertEqual(marker_row.cells[0].text, "{{item_no}}")
            self.assertEqual(marker_row.cells[1].text, "{{item_name}}")
            self.assertEqual(marker_row.cells[3].text, "{{item_quantity}}")
            self.assertEqual(marker_row.cells[4].text, "{{item_unit_price}} ₽")
            self.assertEqual(marker_row.cells[5].text, "{{item_line_total}} ₽")

            for part in ["word/media/image1.png", "word/embeddings/oleObject1.bin"]:
                self.assertEqual(part_hash(output, part), part_hash(SOURCE, part))

    def test_final_render_changes_only_document_xml(self):
        with tempfile.TemporaryDirectory() as tmp:
            prepared = Path(tmp) / "prepared.docx"
            output = Path(tmp) / "output.docx"
            prepare_kartas_template_from_source(SOURCE, prepared)
            _render_kartas_ooxml(self.sample_proposal(), prepared, output)

            doc = Document(output)
            self.assertIn("ООО «ПС Парт»", "\n".join(paragraph.text for paragraph in doc.paragraphs))
            self.assertIn("Сервер KCG Rack 2U", doc.tables[1].rows[1].cells[1].text)
            with ZipFile(prepared) as before, ZipFile(output) as after:
                before_parts = {info.filename: hashlib.sha256(before.read(info.filename)).hexdigest() for info in before.infolist()}
                after_parts = {info.filename: hashlib.sha256(after.read(info.filename)).hexdigest() for info in after.infolist()}
            changed = {name for name in before_parts if before_parts[name] != after_parts[name]}
            self.assertEqual(changed, {"word/document.xml"})


if __name__ == "__main__":
    unittest.main()
